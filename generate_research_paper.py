import openai
from docx import Document

# Set up your OpenAI API key
openai.api_key = 'your-api-key'

# Define the prompts for each section
prompts = {
    "title": "Generate a title for a research paper on using behavioral economics to reform the U.S. healthcare system.",
    "abstract": "Generate an abstract for a research paper titled 'Comprehensive Healthcare for America: Using the Insights of Behavioral Economics to Transform the U.S. Healthcare System'.",
    "introduction": "Generate the introduction section for a research paper titled 'Comprehensive Healthcare for America: Using the Insights of Behavioral Economics to Transform the U.S. Healthcare System'.",
    "section_1": "Generate the section '1. Lessons of Behavioral Economics: Achieving Acceptance of CHA' for a research paper on using behavioral economics to reform the U.S. healthcare system.",
    "section_2": "Generate the section '2. Proposals to Reform the Healthcare System' for a research paper on using behavioral economics to reform the U.S. healthcare system.",
    "section_3": "Generate the section '3. The Tenets of Comprehensive Healthcare for America' for a research paper on using behavioral economics to reform the U.S. healthcare system.",
    "section_4": "Generate the section '4. Healthcare Expenditures under CHA: Potential for Increased Costs Offset by Multiple Opportunities for Savings' for a research paper on using behavioral economics to reform the U.S. healthcare system.",
    "section_5": "Generate the section '5. Overcoming the Political Barriers to CHA' for a research paper on using behavioral economics to reform the U.S. healthcare system.",
    "section_6": "Generate the section '6. The Implementation Process' for a research paper on using behavioral economics to reform the U.S. healthcare system.",
    "section_7": "Generate the section '7. Issues to be Addressed after Implementation' for a research paper on using behavioral economics to reform the U.S. healthcare system.",
    "conclusion": "Generate the conclusion section for a research paper titled 'Comprehensive Healthcare for America: Using the Insights of Behavioral Economics to Transform the U.S. Healthcare System'.",
    "references": "Generate a list of references for a research paper on using behavioral economics to reform the U.S. healthcare system."
}

# Function to generate text from a prompt
def generate_text(prompt):
    response = openai.Completion.create(
      engine="text-davinci-003",
      prompt=prompt,
      max_tokens=1500,
      n=1,
      stop=None,
      temperature=0.7,
    )
    return response.choices[0].text.strip()

# Generate each section of the paper
sections = {}
for section, prompt in prompts.items():
    print(f"Generating {section}...")
    sections[section] = generate_text(prompt)
    print(f"{section} generated.")

# Extract references from the generated references section
references_text = sections['references']
references_list = references_text.split("\n")
references_dict = {str(i + 1): ref for i, ref in enumerate(references_list)}

# Replace placeholder references in sections with numbered references
for number, reference in references_dict.items():
    for section in sections:
        sections[section] = sections[section].replace(f"[{number}]", f"[{number}]")

# Save the generated sections to a Word document
doc = Document()
doc.add_heading(sections["title"], level=1)
doc.add_heading('Abstract', level=2)
doc.add_paragraph(sections["abstract"])
doc.add_heading('Introduction', level=2)
doc.add_paragraph(sections["introduction"])

# Add other sections
for section in ["section_1", "section_2", "section_3", "section_4", "section_5", "section_6", "section_7"]:
    section_title = prompts[section].split("'")[1]
    doc.add_heading(section_title, level=2)
    doc.add_paragraph(sections[section])

doc.add_heading('Conclusion', level=2)
doc.add_paragraph(sections["conclusion"])
doc.add_heading('References', level=2)

for number, reference in references_dict.items():
    doc.add_paragraph(f"[{number}] {reference}")

# Save the document
doc.save("Comprehensive_Healthcare_for_America.docx")

print("Document saved as 'Comprehensive_Healthcare_for_America.docx'.")
